#!/usr/bin/env python3
"""Apply NBSP rules to DOCX after main build (Russian typography for numbers + units).

Inserts non-breaking spaces between numbers and units / abbreviations to satisfy
russian typographical conventions. Skips runs whose font is monospace (code).
"""
from __future__ import annotations
import re
import sys
from docx import Document

NBSP = ' '

# Order matters: longer / more specific patterns first.
PATTERNS = [
    # п. п. (процентный пункт): "п. п." → "п. п."
    (re.compile(r'(\bп\.) +(\bп\.)'), rf'\1{NBSP}\2'),
    # 5 % / 100 % — number + space + %
    (re.compile(r'(\d) +(%)'), rf'\1{NBSP}\2'),
    # 1 кг / 300 г / 5 мл / 10 л / 100 м / 5 см / 30 мм / 2 шт / 25 °C
    (re.compile(r'(\d) +(кг|г|мл|л|см|мм|шт|°C|м)\b'), rf'\1{NBSP}\2'),
    # № 1
    (re.compile(r'№ +(\d)'), rf'№{NBSP}\1'),
    # гл. 1, рис. 3.2, табл. 4.1, с. 12 — abbreviation + number
    (re.compile(r'\b(гл|рис|табл|с)\. +(\d)'), rf'\1.{NBSP}\2'),
    # 2024 г. (год): four-digit year + " г."
    (re.compile(r'(\b\d{4}) +(г\.)'), rf'\1{NBSP}\2'),
    # N × N (e.g. 5 × 4): NBSP on both sides of multiplication sign
    (re.compile(r'(\d) +(×) +(\d)'), rf'\1{NBSP}\2{NBSP}\3'),
]


def is_monospace(run) -> bool:
    name = run.font.name or ''
    return name in {'Courier New', 'Consolas', 'Menlo', 'Monaco', 'monospace'}


def fix_run(run) -> bool:
    if is_monospace(run):
        return False
    txt = run.text
    new = txt
    for pat, repl in PATTERNS:
        new = pat.sub(repl, new)
    if new != txt:
        run.text = new
        return True
    return False


def main(path: str) -> int:
    doc = Document(path)
    n_changes = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if fix_run(r):
                n_changes += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if fix_run(r):
                            n_changes += 1
    doc.save(path)
    print(f'typography_docx: {n_changes} runs modified in {path}')
    return 0


if __name__ == '__main__':
    sys.exit(main(sys.argv[1]))
