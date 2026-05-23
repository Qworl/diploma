#!/usr/bin/env python3
"""Regression check for thesis golden anchors (numbers + strings).

Parses GOLDEN_NUMBERS.md and GOLDEN_STRINGS.md, then checks that every
declared anchor appears at least in one of two sources:
  - the final DOCX (docs/thesis/VKR_Frolov_2026.docx) via python-docx;
  - the thesis notebook (docs/thesis/00_thesis_main.ipynb) via nbformat.

Whitespace is normalised before comparison (NBSP, tab, multiple spaces are
collapsed to a single regular space). Comparison is case-sensitive.

Exit codes:
  0 — every golden anchor is found OR no golden items are configured;
  1 — at least one golden anchor is missing;
  2 — both DOCX and notebook are absent (cannot verify);
  3 — required Python dependency (python-docx / nbformat) not installed.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Iterable

# Default paths are resolved relative to the repository root.
# The script lives in <repo>/docs/thesis/docx_build/, so repo root is parents[3].
_REPO_ROOT = Path(__file__).resolve().parents[3]
_DEFAULT_DOCX = _REPO_ROOT / "docs" / "thesis" / "VKR_Frolov_2026.docx"
_DEFAULT_NOTEBOOK = _REPO_ROOT / "docs" / "thesis" / "00_thesis_main.ipynb"
_DEFAULT_NUMBERS = _REPO_ROOT / "docs" / "thesis" / "GOLDEN_NUMBERS.md"
_DEFAULT_STRINGS = _REPO_ROOT / "docs" / "thesis" / "GOLDEN_STRINGS.md"


# ---------------------------------------------------------------------------
# Normalisation
# ---------------------------------------------------------------------------

_WHITESPACE_RE = re.compile(r"\s+")


def normalize(s: str) -> str:
    """Collapse any whitespace (incl. NBSP \xa0, tab, newlines) to a single space.

    NBSP is part of ``\\s`` for the ``re`` module in Python 3, so a plain
    ``\\s+`` substitution already takes care of it.
    """
    return _WHITESPACE_RE.sub(" ", s).strip()


# ---------------------------------------------------------------------------
# Golden file parsers
# ---------------------------------------------------------------------------

_BACKTICK_RE = re.compile(r"`([^`\n]+)`")


def _iter_sections(lines: Iterable[str]) -> Iterable[tuple[str, list[str]]]:
    """Yield (header_text, body_lines) pairs for ``## ``-level sections.

    Lines before the first ``## `` header are yielded under header ``""``.
    """
    current_header = ""
    current_body: list[str] = []
    for line in lines:
        if line.startswith("## "):
            yield current_header, current_body
            current_header = line[3:].rstrip("\n").strip()
            current_body = []
        else:
            current_body.append(line)
    yield current_header, current_body


def parse_golden_numbers(path: Path) -> list[str]:
    """Parse GOLDEN_NUMBERS.md and return the list of golden numbers.

    The file contains a markdown table with columns
    ``Число | Файл | Раздел/Контекст | Происхождение``. The first column
    holds the golden number wrapped in backticks. Section ``## Пример формата``
    (and any other section whose header contains the word "Пример") is
    skipped entirely — it is illustrative scaffolding, not real data.
    """
    if not path.exists():
        return []
    text = path.read_text(encoding="utf-8")
    numbers: list[str] = []
    for header, body in _iter_sections(text.splitlines(keepends=True)):
        if "Пример" in header:
            continue
        for raw_line in body:
            line = raw_line.rstrip("\n")
            # Only table rows start with a pipe and contain at least one pipe inside.
            if not line.lstrip().startswith("|"):
                continue
            # Skip the table header separator (|---|---|...).
            if set(line.replace("|", "").strip()) <= {"-", ":", " "}:
                continue
            # First cell is between the first two pipes.
            parts = [p.strip() for p in line.strip().strip("|").split("|")]
            if not parts:
                continue
            first = parts[0]
            # Skip the column-name row "Число".
            if first.lower() in {"число", ""}:
                continue
            match = _BACKTICK_RE.search(first)
            if match:
                numbers.append(match.group(1))
    return numbers


_NUMBERED_SECTION_RE = re.compile(r"^\d+\.\s")


def parse_golden_strings(path: Path) -> list[str]:
    """Parse GOLDEN_STRINGS.md and return the list of golden strings.

    File layout — sections ``## 1. ...``, ``## 2. ...``, etc., each filled
    with markdown checkboxes ``- [ ] `value```. Comment-lines without
    backticks (``- (дополнить...)``) are ignored. Only numbered sections
    (header starts with ``<digit>. ``) are considered: this skips ``Правила``,
    ``Чек-листы`` and example/scaffolding sections. Sections whose header
    contains the word "Пример" are likewise skipped.
    """
    if not path.exists():
        return []
    text = path.read_text(encoding="utf-8")
    strings: list[str] = []
    for header, body in _iter_sections(text.splitlines(keepends=True)):
        if "Пример" in header:
            continue
        if not _NUMBERED_SECTION_RE.match(header):
            # Skip non-data sections (Правила, Чек-листы, intro prose, ...).
            continue
        for raw_line in body:
            line = raw_line.rstrip("\n")
            # Process only checkbox bullets to avoid pulling backticks
            # from the prose / "Правила" sections.
            stripped = line.lstrip()
            if not (stripped.startswith("- [ ]") or stripped.startswith("- [x]")):
                continue
            # Strip the leading "- [ ] " / "- [x] " marker, then look at the
            # content. Parenthetical comment items like
            # "- [ ] (наполнить ... `06_references.md` ...)" are placeholders,
            # not actual anchors; skip them.
            content = stripped[5:].lstrip()  # drop "- [ ]" or "- [x]"
            if content.startswith("("):
                continue
            match = _BACKTICK_RE.search(line)
            if match:
                strings.append(match.group(1))
    return strings


# ---------------------------------------------------------------------------
# Source extractors
# ---------------------------------------------------------------------------

def extract_docx_text(path: Path) -> str:
    """Return all text from DOCX paragraphs and table cells, normalised."""
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx is required. Install via: pip install python-docx nbformat"
        ) from exc

    doc = docx.Document(str(path))
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return normalize("\n".join(chunks))


def extract_notebook_text(path: Path) -> str:
    """Return all text from notebook markdown and code cells, normalised."""
    try:
        import nbformat  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "nbformat is required. Install via: pip install python-docx nbformat"
        ) from exc

    nb = nbformat.read(str(path), as_version=4)
    chunks: list[str] = []
    for cell in nb.cells:
        if cell.cell_type in ("markdown", "code"):
            source = cell.source
            if isinstance(source, list):
                source = "".join(source)
            if source:
                chunks.append(source)
    return normalize("\n".join(chunks))


# ---------------------------------------------------------------------------
# Checker
# ---------------------------------------------------------------------------

def check_anchors(
    anchors: list[str],
    sources: dict[str, str],
    kind: str,
    quiet: bool,
) -> list[str]:
    """Check that every anchor occurs in at least one source.

    Parameters
    ----------
    anchors : list of golden anchors (already in display form, with NBSP etc).
    sources : mapping ``source_label -> normalised text``. Empty strings are
        treated as "source missing".
    kind : human-readable category ("golden number" / "golden string"), used
        only in failure messages.
    quiet : when True, do not print per-anchor success lines.

    Returns the list of failure messages (empty on full success).
    """
    failures: list[str] = []
    active_sources = {label: text for label, text in sources.items() if text}
    if not active_sources:
        # Caller is responsible for emitting a higher-level diagnostic.
        return failures
    for anchor in anchors:
        needle = normalize(anchor)
        if not needle:
            continue
        found_in: list[str] = []
        for label, text in active_sources.items():
            if needle in text:
                found_in.append(label)
        if not found_in:
            srcs = " and ".join(active_sources.keys())
            failures.append(
                f"[FAIL] {kind} '{anchor}' not found in {srcs}"
            )
        elif not quiet:
            print(f"[ OK ] {kind} '{anchor}' found in {', '.join(found_in)}")
    return failures


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Regression check for thesis golden numbers and golden strings."
        )
    )
    parser.add_argument(
        "--docx",
        type=Path,
        default=_DEFAULT_DOCX,
        help=f"Path to compiled DOCX (default: {_DEFAULT_DOCX})",
    )
    parser.add_argument(
        "--notebook",
        type=Path,
        default=_DEFAULT_NOTEBOOK,
        help=f"Path to thesis notebook (default: {_DEFAULT_NOTEBOOK})",
    )
    parser.add_argument(
        "--numbers",
        type=Path,
        default=_DEFAULT_NUMBERS,
        help=f"Path to GOLDEN_NUMBERS.md (default: {_DEFAULT_NUMBERS})",
    )
    parser.add_argument(
        "--strings",
        type=Path,
        default=_DEFAULT_STRINGS,
        help=f"Path to GOLDEN_STRINGS.md (default: {_DEFAULT_STRINGS})",
    )
    parser.add_argument(
        "--quiet",
        action="store_true",
        help="Print only the summary line (CI-friendly).",
    )
    args = parser.parse_args()

    # --- Load golden anchors ----------------------------------------------
    numbers = parse_golden_numbers(args.numbers)
    strings = parse_golden_strings(args.strings)
    total_anchors = len(numbers) + len(strings)

    if total_anchors == 0:
        print("[INFO] no golden items configured (GOLDEN_NUMBERS.md / "
              "GOLDEN_STRINGS.md are empty or contain only the example "
              "section). Exiting with success.")
        return 0

    # --- Load sources -----------------------------------------------------
    try:
        if args.docx.exists():
            docx_text = extract_docx_text(args.docx)
        else:
            print(f"[WARN] docx not found at {args.docx}, skipping docx check")
            docx_text = ""
        if args.notebook.exists():
            notebook_text = extract_notebook_text(args.notebook)
        else:
            print(f"[WARN] notebook not found at {args.notebook}, "
                  f"skipping notebook check")
            notebook_text = ""
    except ImportError as exc:
        print(f"[ERROR] {exc}")
        return 3

    if not docx_text and not notebook_text:
        print("[ERROR] neither DOCX nor notebook is available; "
              "cannot verify golden anchors.")
        return 2

    sources = {
        str(args.docx): docx_text,
        str(args.notebook): notebook_text,
    }

    # --- Run the checks ---------------------------------------------------
    failures: list[str] = []
    failures.extend(
        check_anchors(numbers, sources, "golden number", args.quiet)
    )
    failures.extend(
        check_anchors(strings, sources, "golden string", args.quiet)
    )

    # --- Summary ----------------------------------------------------------
    if failures:
        for msg in failures:
            print(msg)
        print(
            f"[SUMMARY] {len(failures)} of {total_anchors} golden anchors "
            f"MISSING ({len(numbers)} numbers + {len(strings)} strings checked)."
        )
        return 1

    print(
        f"[SUMMARY] all {total_anchors} golden anchors found "
        f"({len(numbers)} numbers + {len(strings)} strings)."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
