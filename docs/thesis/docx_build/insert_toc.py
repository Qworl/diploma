#!/usr/bin/env python3
r"""Insert Word TOC field after the abstract/definitions, before chapter 1.

Word TOC (Table of Contents) field syntax: { TOC \o "1-3" \h \z \u }
  \o "1-3" — include Heading 1..3
  \h       — make entries hyperlinks
  \z       — hide tab leader/page numbers in Web layout
  \u       — use applied paragraph outline level

Note: TOC field is updated by Word on first open via prompt or F9.
Until then the placeholder text is shown.

Usage:
    python insert_toc.py [docx_path]
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _make_toc_paragraph(doc):
    """Создаёт параграф с полем Word TOC и возвращает его OOXML-элемент."""
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Оглавление обновляется в Word: ПКМ по полю > Обновить поле, "
        "либо клавиша F9."
    )

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(placeholder)
    r.append(fldChar3)
    return p


def _is_introduction_heading(p) -> bool:
    """Проверка: данный параграф — заголовок «ВВЕДЕНИЕ» (стиль Heading 1)."""
    try:
        style_name = p.style.name
    except Exception:
        return False
    if not style_name.startswith("Heading 1"):
        return False
    text = p.text.strip().upper()
    return text.startswith("ВВЕДЕНИЕ")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", nargs="?", default="VKR_Frolov_2026.docx")
    args = parser.parse_args()

    path = Path(args.docx)
    if not path.exists():
        print(f"ERROR: {path} не найден", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))

    # Идемпотентность: если поле TOC уже есть — пропустить.
    body = doc.element.body
    for instr in body.iter(qn("w:instrText")):
        if instr.text and instr.text.strip().upper().startswith("TOC"):
            print("TOC field already present, skipping insertion")
            return 0

    insert_before = None
    for p in doc.paragraphs:
        if _is_introduction_heading(p):
            insert_before = p
            break

    if insert_before is None:
        print(
            "WARN: anchor heading (ВВЕДЕНИЕ) not found, skipping TOC insertion",
            file=sys.stderr,
        )
        return 1

    # Создаём временно в конце документа заголовок «СОДЕРЖАНИЕ» и TOC-параграф,
    # затем переносим их перед якорным «ВВЕДЕНИЕ».
    heading_p = doc.add_paragraph("СОДЕРЖАНИЕ", style="Heading 1")
    toc_p = _make_toc_paragraph(doc)

    heading_el = heading_p._element
    toc_el = toc_p._element
    body.remove(heading_el)
    body.remove(toc_el)

    insert_before._element.addprevious(heading_el)
    insert_before._element.addprevious(toc_el)

    doc.save(str(path))
    print("TOC inserted before ВВЕДЕНИЕ")
    return 0


if __name__ == "__main__":
    sys.exit(main())
